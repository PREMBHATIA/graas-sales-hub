"""Render a Prospect Brief from a structured dict into DOCX (for Drive upload)
or HTML (for inline preview).

Both renderers consume the same `BriefData` shape so the LLM only has to
produce one format (JSON) — and the on-screen preview matches what the
Google Doc will look like.

The DOCX path uses python-docx and sets explicit table column widths,
margins, font sizes, and paragraph spacing — these survive Google Drive's
DOCX → Doc conversion, which is the whole reason we moved off HTML.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GRAAS_BLUE = RGBColor(0x27, 0x42, 0xFF)
LIGHT_BLUE = "EEF1FF"  # table header fill
YELLOW = "FFF4B8"  # post-call highlight — rows changed by the latest call
GREY = RGBColor(0x66, 0x66, 0x66)
CALLOUT_FILL = "FFF4E5"


# ──────────────────────────────────────────────────────────────────────────────
# DOCX low-level helpers
# ──────────────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_fill: str) -> None:
    """Apply a background fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40,
                      left: int = 80, right: int = 80) -> None:
    """Tighter cell padding than the DOCX default (units = twentieths of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_sub(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY


def _add_status(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True


def _add_para(doc: Document, text: str, size: float = 10.0, italic: bool = False) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.font.italic = True


def _add_kv_para(doc: Document, label: str, value: str, size: float = 10.0) -> None:
    """A 'Label: value' line where label is bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    lbl = p.add_run(f"{label}: ")
    lbl.font.size = Pt(size)
    lbl.font.bold = True
    val = p.add_run(value)
    val.font.size = Pt(size)


def _add_bullets(doc: Document, items: list, size: float = 10.0) -> None:
    for item in items:
        if not item:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(str(item))
        run.font.size = Pt(size)


def _add_table(
    doc: Document,
    headers: list,
    rows: list,
    col_widths_cm: list,
    header_size: float = 9.5,
    cell_size: float = 9.5,
    col_styles: dict = None,
    highlighted_rows: set = None,
) -> None:
    """Build a table with explicit column widths and tighter cell padding.

    rows: list[list[str]] — must match len(headers).
    col_widths_cm: list[float] — column widths in cm, must match len(headers).
    col_styles: optional dict of {col_index: {"size": float, "italic": bool, "color": RGBColor}}
                to override per-column font size / italic / color in DATA cells only.
    highlighted_rows: optional set of 0-based row indices to paint with YELLOW
                fill (post-call change-highlight — surfaces which rows the
                latest call updated/added).
    """
    highlighted_rows = highlighted_rows or set()
    col_styles = col_styles or {}
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"

    # Force fixed table layout — without this OOXML flag, Word and Google Docs
    # auto-fit columns based on cell content length, ignoring our cm widths.
    # Long-content cells then squish other columns, producing "scrambled"
    # tables. With type="fixed", widths are respected.
    tbl_pr = table._tbl.tblPr
    layout_el = OxmlElement("w:tblLayout")
    layout_el.set(qn("w:type"), "fixed")
    tbl_pr.append(layout_el)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(sum(col_widths_cm) * 567)))  # 567 twips ≈ 1cm
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    # Override <w:tblGrid> with our actual column widths in twips. python-docx
    # populates this with equal page-width-divided-by-col-count regardless of
    # what we set on individual cells — and Google Docs uses tblGrid (not tcW)
    # when laying out tables, so the rendered widths come out wrong without
    # this. This is what was making Exec Summary (3-col) render narrower than
    # the stat band (5-col) even though both tblW values were the same.
    _grid = table._tbl.find(qn("w:tblGrid"))
    if _grid is not None:
        for _child in list(_grid):
            _grid.remove(_child)
        for _w in col_widths_cm:
            _gc = OxmlElement("w:gridCol")
            _gc.set(qn("w:w"), str(int(_w * 567)))
            _grid.append(_gc)

    # Set widths on every cell of every row (DOCX needs this redundantly)
    for col_idx, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(w)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_cell_shading(cell, LIGHT_BLUE)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(h))
        run.font.size = Pt(header_size)
        run.font.bold = True

    # Data rows
    for r_idx, row in enumerate(rows, start=1):
        # 0-based data-row index for highlight check (r_idx is table-row index
        # which counts the header at 0)
        data_idx = r_idx - 1
        is_highlighted = data_idx in highlighted_rows
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            if is_highlighted:
                _set_cell_shading(cell, YELLOW)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(val) if val is not None else "")
            style = col_styles.get(c_idx, {})
            run.font.size = Pt(style.get("size", cell_size))
            if style.get("italic"):
                run.font.italic = True
            if style.get("color"):
                run.font.color.rgb = style["color"]

    # Tiny trailing spacer paragraph — without this, Word and Google Docs
    # visually merge consecutive <w:tbl> blocks into one continuous table
    # with no border break (Exec Summary rows fused with the stat band).
    # A 4pt-font empty paragraph forces a clean disconnect between tables
    # without adding noticeable vertical space.
    _spacer = doc.add_paragraph()
    _spacer.paragraph_format.space_before = Pt(0)
    _spacer.paragraph_format.space_after = Pt(0)
    _spacer_run = _spacer.add_run("")
    _spacer_run.font.size = Pt(4)


def _add_callout_box(doc: Document, lines: list) -> None:
    """Single-cell amber-tinted callout box for Conflicts & Unknowns."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.width = Cm(19.5)
    _set_cell_shading(cell, CALLOUT_FILL)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell.text = ""
    for i, (label, val) in enumerate(lines):
        if not val:
            continue
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        lbl = p.add_run(f"{label}: ")
        lbl.font.size = Pt(9.5)
        lbl.font.bold = True
        v = p.add_run(str(val))
        v.font.size = Pt(9.5)



# ──────────────────────────────────────────────────────────────────────────────
# Hyperlink helper (LinkedIn names in key_people)
# ──────────────────────────────────────────────────────────────────────────────
def _add_hyperlink(paragraph, url: str, text: str, *, bold: bool = True,
                   size: float = 10.0, color: str = "2742FF") -> None:
    """Append a real clickable hyperlink run to an existing paragraph.

    python-docx has no native hyperlink API, so we build the OOXML by hand and
    register the external relationship. Survives the Drive DOCX→Doc conversion.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _clean_linkedin(url) -> str:
    """Return a usable https LinkedIn URL, or '' if it's not a real one."""
    s = ("" if url is None else str(url)).strip()
    if not s:
        return ""
    low = s.lower()
    if "linkedin.com/in/" not in low and "linkedin.com/pub/" not in low:
        return ""
    if not low.startswith("http"):
        s = "https://" + s.lstrip("/")
    return s


def _as_list(v) -> list:
    """Coerce a string-or-list field into a clean list of non-empty items."""
    if v is None:
        return []
    if isinstance(v, str):
        v = [v]
    return [x for x in v if x]


# ──────────────────────────────────────────────────────────────────────────────
# Public renderer — DOCX (two-pager)
# ──────────────────────────────────────────────────────────────────────────────
def render_brief_docx(data: dict) -> bytes:
    """Render a two-pager brief data dict into DOCX bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    company = data.get("company") or "[Company]"
    header = data.get("header", {}) or {}

    # ── Header ────────────────────────────────────────────────────────────────
    _add_h1(doc, f"Prospect Brief — {company}")
    sub_parts = ["Graas Pre-Sales", "Confidential"]
    if header.get("date_prepared"):
        sub_parts.append(f"Prepared {header['date_prepared']}")
    if header.get("meeting_date"):
        sub_parts.append(f"Meeting {header['meeting_date']}")
    if header.get("market"):
        sub_parts.append(f"Market {header['market']}")
    _add_sub(doc, " · ".join(sub_parts))
    if header.get("meeting_context"):
        _add_para(doc, header["meeting_context"], size=9.0, italic=True)

    # ── Snapshot (100% factual header strip) ──────────────────────────────────
    sb = data.get("summary_boxes") or {}
    if isinstance(sb, dict) and any(sb.values()):
        _add_h2(doc, "Snapshot")
        _add_table(
            doc,
            ["Industry", "Type", "Revenue"],
            [[sb.get("industry", ""), sb.get("type", ""), sb.get("revenue", "")]],
            [6.0, 5.0, 7.0],
        )
        _add_table(
            doc,
            ["Comps", "Scale"],
            [[sb.get("comps", ""), sb.get("scale", "")]],
            [9.0, 9.0],
        )

    # ── Who's who (the buying group) ──────────────────────────────────────────
    people = data.get("key_people") or []
    if people:
        _add_h2(doc, "Who's who")
        for person in people:
            if not isinstance(person, dict):
                continue
            name = (person.get("name") or "").strip()
            desig = (person.get("designation") or "").strip()
            if not name and not desig:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(0)
            li = _clean_linkedin(person.get("linkedin"))
            if name and li:
                _add_hyperlink(p, li, name, bold=True, size=10.0)
            elif name:
                r = p.add_run(name)
                r.font.bold = True
                r.font.size = Pt(10)
            if desig:
                r2 = p.add_run((" — " if name else "") + desig)
                r2.font.size = Pt(10)
            play = (person.get("play") or "").strip()
            if play:
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(2)
                pp.paragraph_format.left_indent = Cm(0.4)
                rp = pp.add_run(play)
                rp.font.size = Pt(9.0)
                rp.font.color.rgb = GREY

    # ── The stack (systems · vendors · owners · verdict) ──────────────────────
    stack = data.get("stack") or []
    if stack:
        _add_h2(doc, "The stack — what they run, who owns it")
        rows = []
        for s in stack:
            if not isinstance(s, dict):
                continue
            src = (s.get("source") or "").strip()
            conf = (s.get("confidence") or "").strip()
            src_conf = " · ".join([x for x in (conf, src) if x])
            rows.append([
                s.get("layer", ""),
                s.get("system", ""),
                s.get("vendor", ""),
                s.get("owner", ""),
                s.get("verdict", ""),
                src_conf,
            ])
        _add_table(
            doc,
            ["Layer", "System", "Vendor / SI", "Owner", "Verdict", "Source · conf"],
            rows,
            [2.2, 3.4, 3.2, 2.8, 3.8, 2.6],
            col_styles={5: {"size": 8.0, "italic": True, "color": GREY}},
        )

    # ── Where Graas might fit (hypotheses, question-framed) ────────────────────
    fit = data.get("graas_fit") or []
    if fit:
        _add_h2(doc, "Where Graas might fit (to test, not assert)")
        rows = []
        for f in fit:
            if not isinstance(f, dict):
                continue
            rows.append([f.get("where", ""), f.get("graas_offering", ""),
                         f.get("fit", ""), f.get("verify", "")])
        _add_table(
            doc,
            ["Where", "Graas offering", "Could Graas fit? (hypothesis)", "Verify in the meeting"],
            rows,
            [2.6, 3.4, 7.0, 5.0],
            col_styles={1: {"size": 9.0, "color": GRAAS_BLUE}},
        )

    # ── Landmines ─────────────────────────────────────────────────────────────
    do_not = _as_list(data.get("do_not"))
    if do_not:
        _add_h2(doc, "Landmines — do NOT")
        _add_bullets(doc, do_not, size=9.5)

    # ── Wedges worth exploring (exploratory, not a recommendation) ────────────
    wedges = _as_list(data.get("wedges_worth_exploring"))
    if wedges:
        _add_h2(doc, "Wedges worth exploring")
        _add_bullets(doc, wedges, size=10.0)

    # ── Appendix ──────────────────────────────────────────────────────────────
    discovery = _as_list(data.get("discovery"))
    appendix_research = [c for c in (data.get("appendix_research") or []) if isinstance(c, dict)]
    if discovery or appendix_research:
        _add_h2(doc, "Appendix")
        if discovery:
            _add_h3(doc, "Discovery questions")
            _add_bullets(doc, [f"{i}. {q}" for i, q in enumerate(discovery, start=1)], size=9.5)
        if appendix_research:
            _add_h3(doc, "Further confirmed research")
            rows = [[c.get("fact", ""), c.get("source", "")] for c in appendix_research]
            _add_table(
                doc,
                ["Fact (confirmed)", "Source"],
                rows,
                [13.0, 5.0],
                col_styles={1: {"size": 8.5, "italic": True, "color": GREY}},
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Public renderer — HTML preview (two-pager)
# ──────────────────────────────────────────────────────────────────────────────
def _esc(s: Any) -> str:
    import html as _h
    return _h.escape("" if s is None else str(s))


def render_brief_html(data: dict) -> str:
    """Render the same two-pager brief data as a tight HTML preview."""
    company = _esc(data.get("company") or "[Company]")
    header = data.get("header", {}) or {}

    parts: list = []
    parts.append("""<!doctype html><html><head><meta charset='utf-8'><style>
body { font-family: Calibri, Arial, sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.25; margin: 0; padding: 0 12px; }
h1 { font-size: 16pt; color: #2742FF; margin: 4pt 0 2pt 0; }
h2 { font-size: 11pt; color: #2742FF; margin: 11pt 0 4pt 0; }
h3 { font-size: 10pt; color: #2742FF; margin: 8pt 0 2pt 0; }
p { margin: 2pt 0; }
ul { margin: 2pt 0 4pt 0; padding-left: 18pt; }
li { margin: 0 0 1pt 0; }
.sub { color: #666; font-size: 8.5pt; margin: 0; }
.ctx { color: #666; font-size: 9pt; font-style: italic; margin: 1pt 0 4pt 0; }
table { border-collapse: collapse; width: 100%; margin: 3pt 0; }
th, td { border: 1px solid #bbb; padding: 2pt 6pt; text-align: left; vertical-align: top; font-size: 9.5pt; line-height: 1.2; }
th { background: #eef1ff; font-weight: bold; }
.person { margin: 2pt 0 0 0; }
.person .name { font-weight: bold; }
.person a { color: #2742FF; text-decoration: underline; font-weight: bold; }
.play { color: #666; font-size: 9pt; margin: 0 0 3pt 12pt; }
td.dim { font-size: 8pt; font-style: italic; color: #777; }
td.offering { color: #2742FF; font-weight: bold; font-size: 9pt; }
.landmine li { color: #7a1f1f; }
.appx { border-top: 2px solid #2742FF; margin-top: 10pt; padding-top: 2pt; }
</style></head><body>""")

    parts.append(f"<h1>Prospect Brief — {company}</h1>")
    sub_parts = ["Graas Pre-Sales", "Confidential"]
    if header.get("date_prepared"):
        sub_parts.append(f"Prepared {_esc(header['date_prepared'])}")
    if header.get("meeting_date"):
        sub_parts.append(f"Meeting {_esc(header['meeting_date'])}")
    if header.get("market"):
        sub_parts.append(f"Market {_esc(header['market'])}")
    parts.append(f"<p class='sub'>{' · '.join(sub_parts)}</p>")
    if header.get("meeting_context"):
        parts.append(f"<p class='ctx'>{_esc(header['meeting_context'])}</p>")

    # Snapshot boxes
    sb = data.get("summary_boxes") or {}
    if isinstance(sb, dict) and any(sb.values()):
        parts.append("<h2>Snapshot</h2>")
        parts.append("<table><tr><th>Industry</th><th>Type</th><th>Revenue</th></tr>")
        parts.append(f"<tr><td>{_esc(sb.get('industry'))}</td><td>{_esc(sb.get('type'))}</td>"
                     f"<td>{_esc(sb.get('revenue'))}</td></tr></table>")
        parts.append("<table><tr><th>Comps</th><th>Scale</th></tr>")
        parts.append(f"<tr><td>{_esc(sb.get('comps'))}</td><td>{_esc(sb.get('scale'))}</td></tr></table>")

    # Who's who
    people = data.get("key_people") or []
    if people:
        parts.append("<h2>Who's who</h2>")
        for person in people:
            if not isinstance(person, dict):
                continue
            name = (person.get("name") or "").strip()
            desig = (person.get("designation") or "").strip()
            if not name and not desig:
                continue
            li = _clean_linkedin(person.get("linkedin"))
            if name and li:
                name_html = f"<a href='{_esc(li)}' target='_blank'>{_esc(name)}</a>"
            else:
                name_html = f"<span class='name'>{_esc(name)}</span>"
            tail = f" — {_esc(desig)}" if desig else ""
            parts.append(f"<p class='person'>{name_html}{tail}</p>")
            play = (person.get("play") or "").strip()
            if play:
                parts.append(f"<p class='play'>{_esc(play)}</p>")

    # The stack
    stack = data.get("stack") or []
    if stack:
        parts.append("<h2>The stack — what they run, who owns it</h2><table>")
        parts.append("<tr><th>Layer</th><th>System</th><th>Vendor / SI</th><th>Owner</th>"
                     "<th>Verdict</th><th>Source · conf</th></tr>")
        for s in stack:
            if not isinstance(s, dict):
                continue
            src = (s.get("source") or "").strip()
            conf = (s.get("confidence") or "").strip()
            src_conf = " · ".join([x for x in (conf, src) if x])
            parts.append(
                f"<tr><td>{_esc(s.get('layer'))}</td><td>{_esc(s.get('system'))}</td>"
                f"<td>{_esc(s.get('vendor'))}</td><td>{_esc(s.get('owner'))}</td>"
                f"<td>{_esc(s.get('verdict'))}</td><td class='dim'>{_esc(src_conf)}</td></tr>"
            )
        parts.append("</table>")

    # Where Graas might fit
    fit = data.get("graas_fit") or []
    if fit:
        parts.append("<h2>Where Graas might fit (to test, not assert)</h2><table>")
        parts.append("<tr><th>Where</th><th>Graas offering</th>"
                     "<th>Could Graas fit? (hypothesis)</th>"
                     "<th>Verify in the meeting</th></tr>")
        for f in fit:
            if not isinstance(f, dict):
                continue
            parts.append(
                f"<tr><td>{_esc(f.get('where'))}</td>"
                f"<td class='offering'>{_esc(f.get('graas_offering'))}</td>"
                f"<td>{_esc(f.get('fit'))}</td>"
                f"<td>{_esc(f.get('verify'))}</td></tr>"
            )
        parts.append("</table>")

    # Landmines
    do_not = _as_list(data.get("do_not"))
    if do_not:
        parts.append("<h2>Landmines — do NOT</h2><ul class='landmine'>")
        for d in do_not:
            parts.append(f"<li>{_esc(d)}</li>")
        parts.append("</ul>")

    # Wedges worth exploring
    wedges = _as_list(data.get("wedges_worth_exploring"))
    if wedges:
        parts.append("<h2>Wedges worth exploring</h2><ul>")
        for w in wedges:
            parts.append(f"<li>{_esc(w)}</li>")
        parts.append("</ul>")

    # Appendix
    discovery = _as_list(data.get("discovery"))
    appendix_research = [c for c in (data.get("appendix_research") or []) if isinstance(c, dict)]
    if discovery or appendix_research:
        parts.append("<h2 class='appx'>Appendix</h2>")
        if discovery:
            parts.append("<h3>Discovery questions</h3><ul>")
            for i, q in enumerate(discovery, start=1):
                parts.append(f"<li>{i}. {_esc(q)}</li>")
            parts.append("</ul>")
        if appendix_research:
            parts.append("<h3>Further confirmed research</h3><table>")
            parts.append("<tr><th>Fact (confirmed)</th><th>Source</th></tr>")
            for c in appendix_research:
                parts.append(
                    f"<tr><td>{_esc(c.get('fact'))}</td>"
                    f"<td class='dim'>{_esc(c.get('source'))}</td></tr>"
                )
            parts.append("</table>")

    parts.append("</body></html>")
    return "".join(parts)
