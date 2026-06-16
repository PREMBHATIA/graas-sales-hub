"""Render an All-e Solution Doc from a structured dict into DOCX (for Drive)
or HTML (for inline preview).

Mirrors services/brief_renderer.py — same toolkit (python-docx, narrow margins,
explicit table widths, OOXML cell padding, Graas blue), different schema.

Solution Doc sits BETWEEN the Prospect Brief and the Proposal:
  Brief (facts) → **Solution (what we build)** → Proposal (commercials)

Four sections:
  1. Core functionality — per-agent: persona, surfaces, what it does, phase
  2. Key agent KPIs — target, baseline, baseline source
  3. Missing fields & data gaps — what to ask the customer for next
  4. Timeline — phase, duration, milestone

Conscious omissions: no pricing, no objection handling, no demo planning —
those belong in Create Proposal (pricing) or Resources (demo material).
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GRAAS_BLUE = RGBColor(0x27, 0x42, 0xFF)
LIGHT_BLUE = "EEF1FF"
GREY = RGBColor(0x66, 0x66, 0x66)
AMBER_FILL = "FFF4E5"


# ──────────────────────────────────────────────────────────────────────────────
# Low-level DOCX helpers (duplicated minimally from brief_renderer to keep this
# module standalone; if a third doc type lands we should factor these out)
# ──────────────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40,
                      left: int = 80, right: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY


def _add_status(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GRAAS_BLUE


def _add_para(doc, text, size=10.0, italic=False):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.font.italic = True


def _add_kv_para(doc, label, value, size=10.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    lbl = p.add_run(f"{label}: ")
    lbl.font.size = Pt(size)
    lbl.font.bold = True
    val = p.add_run(value)
    val.font.size = Pt(size)


def _add_table(doc, headers, rows, col_widths_cm,
               header_size=9.5, cell_size=9.5, col_styles=None):
    col_styles = col_styles or {}
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"

    for col_idx, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(w)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_cell_shading(cell, LIGHT_BLUE)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(h))
        run.font.size = Pt(header_size)
        run.font.bold = True

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(val) if val is not None else "")
            style = col_styles.get(c_idx, {})
            run.font.size = Pt(style.get("size", cell_size))
            if style.get("italic"):
                run.font.italic = True
            if style.get("color"):
                run.font.color.rgb = style["color"]


# ──────────────────────────────────────────────────────────────────────────────
# Public DOCX renderer
# ──────────────────────────────────────────────────────────────────────────────

def render_soln_docx(data: dict) -> bytes:
    """Render a solution dict into DOCX bytes."""
    doc = Document()

    # Narrow margins so the wide tables breathe
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    company = data.get("company") or "[Company]"
    header = data.get("header", {}) or {}

    # ── Header ───────────────────────────────────────────────────────────────
    _add_h1(doc, f"Solution Architecture — {company}")
    sub_parts = ["Graas Solutioning", "Confidential"]
    if header.get("date_prepared"):
        sub_parts.append(f"Prepared {header['date_prepared']}")
    if header.get("based_on_brief"):
        sub_parts.append(f"Based on brief: {header['based_on_brief']}")
    _add_sub(doc, " · ".join(sub_parts))
    if header.get("status"):
        _add_status(doc, f"Status: {header['status']}")

    # ── Executive Summary (one-liner pitch of the proposed solution) ─────────
    if data.get("executive_summary"):
        _add_h2(doc, "Executive Summary")
        _add_para(doc, data["executive_summary"], size=10)

    # ── 1. Core functionality ─────────────────────────────────────────────────
    core = data.get("core_functionality") or []
    if core:
        _add_h2(doc, "1. Core functionality")
        rows = []
        for r in core:
            surfaces = r.get("surfaces", [])
            if isinstance(surfaces, list):
                surfaces = " / ".join(surfaces)
            rows.append([
                r.get("agent_name", ""),
                r.get("persona", ""),
                surfaces or "",
                r.get("what_it_does", ""),
                r.get("phase", ""),
            ])
        _add_table(
            doc,
            headers=["Agent", "Persona", "Surfaces", "What it does", "Phase"],
            rows=rows,
            col_widths_cm=[3.5, 2.8, 2.8, 8.4, 2.0],
        )

    # ── 2. Key agent KPIs ────────────────────────────────────────────────────
    kpis = data.get("key_agent_kpis") or []
    if kpis:
        _add_h2(doc, "2. Key agent KPIs")
        rows = []
        for k in kpis:
            rows.append([
                k.get("agent", ""),
                k.get("kpi", ""),
                k.get("target", ""),
                k.get("baseline", ""),
                k.get("baseline_source", ""),
            ])
        _add_table(
            doc,
            headers=["Agent", "KPI", "Target", "Baseline", "Baseline source"],
            rows=rows,
            col_widths_cm=[3.0, 5.5, 2.8, 3.4, 4.8],
            # Baseline-source styled as a footnote
            col_styles={4: {"size": 6.5, "italic": True, "color": GREY}},
        )

    # ── 3. Missing fields & data gaps ────────────────────────────────────────
    missing = data.get("missing_fields") or []
    if missing:
        _add_h2(doc, "3. Missing fields & data gaps")
        rows = []
        for m in missing:
            rows.append([
                m.get("field", ""),
                m.get("why_needed", ""),
                m.get("owner", ""),
                m.get("ask", ""),
            ])
        _add_table(
            doc,
            headers=["Field / data needed", "Why needed", "Owner", "Ask"],
            rows=rows,
            col_widths_cm=[4.5, 5.5, 3.0, 6.5],
        )
        _add_sub(doc, "(blocking — confirm before final commercials)")

    # ── 4. Timeline ──────────────────────────────────────────────────────────
    timeline = data.get("timeline") or []
    if timeline:
        _add_h2(doc, "4. Timeline")
        rows = []
        for t in timeline:
            rows.append([
                t.get("phase", ""),
                t.get("duration", ""),
                t.get("milestone", ""),
            ])
        _add_table(
            doc,
            headers=["Phase", "Duration", "Milestone"],
            rows=rows,
            col_widths_cm=[5.0, 3.0, 11.5],
        )

    # ── Appendix: reference patterns consulted (if any) ──────────────────────
    refs = data.get("reference_patterns") or []
    if refs:
        _add_h2(doc, "Appendix: Reference patterns consulted")
        _add_para(doc, "  ·  ".join(refs), size=8.5, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Public HTML renderer (for inline preview only)
# ──────────────────────────────────────────────────────────────────────────────

def _esc(s: Any) -> str:
    import html as _h
    return _h.escape("" if s is None else str(s))


def render_soln_html(data: dict) -> str:
    company = _esc(data.get("company") or "[Company]")
    header = data.get("header", {}) or {}

    parts: list = []
    parts.append("""<!doctype html><html><head><meta charset='utf-8'><style>
body { font-family: Calibri, Arial, sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.25; margin: 0; padding: 0 12px; }
h1 { font-size: 16pt; color: #2742FF; margin: 4pt 0 2pt 0; }
h2 { font-size: 11pt; color: #2742FF; margin: 10pt 0 4pt 0; }
.sub { color: #666; font-size: 8.5pt; margin: 0; }
.status { color: #2742FF; font-size: 9pt; font-weight: bold; margin: 2pt 0 6pt 0; }
table { border-collapse: collapse; width: 100%; margin: 3pt 0; }
th, td { border: 1px solid #bbb; padding: 2pt 6pt; text-align: left; vertical-align: top; font-size: 9.5pt; line-height: 1.2; }
th { background: #eef1ff; font-weight: bold; }
td.src { font-size: 7pt; font-style: italic; color: #777; }
.appendix { color: #777; font-size: 8.5pt; font-style: italic; }
</style></head><body>""")

    parts.append(f"<h1>Solution Architecture — {company}</h1>")
    sub_parts = ["Graas Solutioning", "Confidential"]
    if header.get("date_prepared"):
        sub_parts.append(f"Prepared {_esc(header['date_prepared'])}")
    if header.get("based_on_brief"):
        sub_parts.append(f"Based on brief: {_esc(header['based_on_brief'])}")
    parts.append(f"<p class='sub'>{' · '.join(sub_parts)}</p>")
    if header.get("status"):
        parts.append(f"<p class='status'>Status: {_esc(header['status'])}</p>")

    if data.get("executive_summary"):
        parts.append("<h2>Executive Summary</h2>")
        parts.append(f"<p>{_esc(data['executive_summary'])}</p>")

    core = data.get("core_functionality") or []
    if core:
        parts.append("<h2>1. Core functionality</h2><table>")
        parts.append("<tr><th>Agent</th><th>Persona</th><th>Surfaces</th><th>What it does</th><th>Phase</th></tr>")
        for r in core:
            surfaces = r.get("surfaces", [])
            if isinstance(surfaces, list):
                surfaces = " / ".join(surfaces)
            parts.append(
                f"<tr><td>{_esc(r.get('agent_name'))}</td>"
                f"<td>{_esc(r.get('persona'))}</td>"
                f"<td>{_esc(surfaces)}</td>"
                f"<td>{_esc(r.get('what_it_does'))}</td>"
                f"<td>{_esc(r.get('phase'))}</td></tr>"
            )
        parts.append("</table>")

    kpis = data.get("key_agent_kpis") or []
    if kpis:
        parts.append("<h2>2. Key agent KPIs</h2><table>")
        parts.append("<tr><th>Agent</th><th>KPI</th><th>Target</th><th>Baseline</th><th>Baseline source</th></tr>")
        for k in kpis:
            parts.append(
                f"<tr><td>{_esc(k.get('agent'))}</td>"
                f"<td>{_esc(k.get('kpi'))}</td>"
                f"<td>{_esc(k.get('target'))}</td>"
                f"<td>{_esc(k.get('baseline'))}</td>"
                f"<td class='src'>{_esc(k.get('baseline_source'))}</td></tr>"
            )
        parts.append("</table>")

    missing = data.get("missing_fields") or []
    if missing:
        parts.append("<h2>3. Missing fields &amp; data gaps</h2><table>")
        parts.append("<tr><th>Field / data needed</th><th>Why needed</th><th>Owner</th><th>Ask</th></tr>")
        for m in missing:
            parts.append(
                f"<tr><td>{_esc(m.get('field'))}</td>"
                f"<td>{_esc(m.get('why_needed'))}</td>"
                f"<td>{_esc(m.get('owner'))}</td>"
                f"<td>{_esc(m.get('ask'))}</td></tr>"
            )
        parts.append("</table>")
        parts.append("<p class='sub'>(blocking — confirm before final commercials)</p>")

    timeline = data.get("timeline") or []
    if timeline:
        parts.append("<h2>4. Timeline</h2><table>")
        parts.append("<tr><th>Phase</th><th>Duration</th><th>Milestone</th></tr>")
        for t in timeline:
            parts.append(
                f"<tr><td>{_esc(t.get('phase'))}</td>"
                f"<td>{_esc(t.get('duration'))}</td>"
                f"<td>{_esc(t.get('milestone'))}</td></tr>"
            )
        parts.append("</table>")

    refs = data.get("reference_patterns") or []
    if refs:
        parts.append("<h2>Appendix: Reference patterns consulted</h2>")
        parts.append(f"<p class='appendix'>{' · '.join(_esc(r) for r in refs)}</p>")

    parts.append("</body></html>")
    return "".join(parts)
